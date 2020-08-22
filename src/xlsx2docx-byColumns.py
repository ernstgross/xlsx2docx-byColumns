#!/usr/bin/env python3

"""xlsx2docx-byColumns.py: Generates from an xlsx formated file (MS-Excel), 
which provides content, styles and replace-strings, several docx files.
"""

__author__     = "Ernst Gross"
__copyright__  = "Copyright 2020"
__license__    = "MIT"
__version__    = "0.1.0"
__maintainer__ = "Ernst Gross"
__email__      = "ernst@grossmusik.de"
__status__     = "Development"
__credits__    = [ "" ]

import os
import sys
import time

import tomlkit
from tomlkit.toml_document import TOMLDocument
from tomlkit.toml_file     import TOMLFile

import openpyxl 
from openpyxl.utils import get_column_letter, column_index_from_string

import docx


def getCurrentTimeAsString():
    """Return the current time as a formatted string. Example: "2020-08-19_151221".
    """
    return time.strftime("%Y-%m-%d_%H%M%S")


def appendTimestampToFilename(filename):
    """Append the current time to a given filename. Example: "bla.txt" => "bla_2020-08-19_151221.txt".
    Returns a string with the new timestamped filename.
    """

    name, ext = os.path.splitext(filename)
    return "{name}_{uid}{ext}".format(name=name, uid=getCurrentTimeAsString(), ext=ext)


def ensureDirectoryExists(filename):
    baseFolder = os.path.dirname(filename)
    os.makedirs(baseFolder, exist_ok=True)


def appendTimestampToFilenames(filenames):
    """Append the current time to a list of given filename. Example: "bla.txt" => "bla_2020-08-19_151221.txt".
    Returns a list of strings with timestamped filenames.
    """

    filenamesWithTimestamp = []
   
    for filename in filenames:
        filenamesWithTimestamp.append(appendTimestampToFilename(filename))
        ensureDirectoryExists(filename)

    return filenamesWithTimestamp

def replaceInParagraph(para, replace, content, style, contextString):
    """Search for replace string in the given paragraph and replace it with the given content and style.

     Parameters
    ----------
    para : docx.Paragraph
        The docx.Paragraph that will be edited.
    replace : string
        If the replace string is found, the complete paragraph is set to the given conten.
    content: string
        This content will be used to set as paragraph where the replace string is found.
    style : string
        Name of the style that will be used for the paragraph. If style is 'None', the original style of the paragraph is unchanged.
    
    Returns
    -------
    para: docx.Paragraph
        The edited docx.Paragraph.
   
    """

    if replace in para.text:
        print("In ", contextString, ": replace_paragraph: '", para.text, "' \tby '", content, "'", sep="")
        para.text = content

        if style and style != "None":
            print("Set style to '", style, "'", sep="")
            para.style = style
    
    return para


def replaceParagraphInDoc(doc, replace, content, style):
    """Search for replace string in all paragraphs in the given document and replace them with the given content and style.

    The current implementation of docx supports paragraphs in:
    1. document
    2. document tables
    3. sections
        a) header
        b) header tables
        b) footer 
        c) footer tables
    
    Parameters
    ----------
    doc : docx.Document
        The docx.Document that will be edited.
    replace : string
        If the replace string is found, the complete paragraph is set to the given conten.
    content: string
        This content will be used to set as paragraph where the replace string is found.
    style : string
        Name of the style that will be used for the paragraph. If style is 'None', the original style of the paragraph is unchanged.
    
    Returns
    -------
    doc: docx.Document
        The edited docx.Document.
    """

    # #############################################
    # Search in normal paragraphs of the document.
    # #############################################
    for para in doc.paragraphs:
        para = replaceInParagraph(para, replace, content, style, contextString="doc.paragraphs")

    # #############################################
    # Search in tables of the document for paragraphs.
    # #############################################
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para = replaceInParagraph(para, replace, content, style, contextString="doc.tables")
    
    # ###################################################################################
    # Search in all sections in their footer and header and their tables for paragraphs.
    # ###################################################################################
    for section in doc.sections:
        for para in section.header.paragraphs:
            para = replaceInParagraph(para, replace, content, style, contextString="section.header")

        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        para = replaceInParagraph(para, replace, content, style, contextString="section.header.tables")

        for para in section.footer.paragraphs:
            para = replaceInParagraph(para, replace, content, style, contextString="section.footer")

        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        para = replaceInParagraph(para, replace, content, style, contextString="section.footer.tables")
    
    return doc


def getUsername():
    """Return current user as a string.
    
    Parameters
    ----------

    Returns
    -------
    username: string
        The username from the environment variables.
    """

    username = "unknown"

    if os.name == "nt":
        username = os.environ["USERNAME"]
    else:
        username = os.environ["USER"]
    
    return username


def generateDocxFromXlsx(sheet, contentStartRow, commandColumn, styleColumn, replaceColumn, contentColumns, templateFilename, outputFilenames):
    """Generate the output docx files from the input xlsx sheet.
    
    Parameters
    ----------
    sheet : openpyxl.sheet
        Input data to generate the output.
    contentStartRow : int (first row is 1)
        Index of xlsx row where the content starts.
    commandColumn : int (first column is 1)
        Index of xlsx column of the commands. Supported commands are 
        1. replace_paragraph - searches for the replace string in all paragraphs of the document and if found the paragraph will be set to what is defined in the content cell.
        2. add_paragraph - add/append a new paragraph.
    styleColumn : int (first column is 1)
        Index of xlsx column of the strings that name the style to be used. 
        If empty or "None", then in case of command 
        1. replace_paragraph - the paragraph style in the template will not be changed, e.g. if replace_paragraph is used.
        2. add_paragraph - the default paragraph style defined in the template will be used.
    replaceColumn : int (first column is 1)
        Index of xlsx column of the search strings to be replaced with the content.
    contentColumns : list of int
        Indexes of xlsx columns of the provided content
    outputTemplateFilename : string
        The template filename
    outputFilenames : list of string
        The filenames of the generated output files.

    Returns
    -------
    nothing
    """
    outputGeneratedFilenamesWithTimestamp = appendTimestampToFilenames(outputFilenames)

    for outputGeneratedFilenameWithTimestamp, contentColumn in zip(outputGeneratedFilenamesWithTimestamp, contentColumns):
        #######################################################################################
        print("\nGenerate timestamped output file: '", outputGeneratedFilenameWithTimestamp, "' based on template: '", templateFilename, "'", sep="")
        #######################################################################################
        doc = docx.Document(templateFilename)

        doc.core_properties.author = getUsername()
        print("author:", doc.core_properties.author)

        for row in range(contentStartRow, sheet.max_row+1, 1):
            command = sheet.cell(row=row, column=commandColumn ).value
            replace = sheet.cell(row=row, column=replaceColumn ).value
            content = sheet.cell(row=row, column=contentColumn ).value
            style   = sheet.cell(row=row, column=styleColumn   ).value

            if command: 
                # print(command, style, content)

                if "replace_paragraph" in command:
                    doc = replaceParagraphInDoc(doc=doc, replace=replace, content=content, style=style)
                elif "add_paragraph" in command:
                    if content:
                        doc.add_paragraph(content, style=style)
                else:
                    print("Unexpected object:", command, "with content:", content)

        #######################################################################################
        # print("Save to new generated MS-Word file:", outputGeneratedFilenameWithTimestamp)
        #######################################################################################
        doc.save(outputGeneratedFilenameWithTimestamp)

##################################################################################################################################################
# The next string for the defaultConfiguration is intentionally placed out of the corresponding function of generateDefaultConfigurationTomlFile()
# Because there should not be any leading indentation as required in a function.
##################################################################################################################################################
defaultConfiguration = """# This is the TOML formated configuration file for xlsx2docx-byColumns.py.
# Please used this autogenerated defaults as an example to configure them to your needs.
 
[sourceData]
filename         = 'sourceDataExample.xlsx'
contentSheetname = "Content"
contentStartRow  = 2
commandColumn    = 1
styleColumn      = 2
replaceColumn    = 3

# "contentColumns" must correspont to the "generatedData.filenames". 
# That means identical number of elements and same sequence.
contentColumns = [
    4, 
    5, 
    6
]

[template]
filename = "src/xlsx2docx-byColumns_TemplateExample.docx"

[generatedData]
# "filenames" must correspont to the "sourceData.contentColumns". 
# That means identical number of elements and same sequence.
filenames = [
  "generatedDataFile1.docx",
  "generatedDataFile2.docx",
  "generatedDataFile3.docx",
]
"""

def generateDefaultConfigurationTomlFile(configurationFilename):
    print("!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!")
    print("!!  Default configuration was written to file '", configurationFilename,"'  !!", sep="") 
    print("!!  It is assumed, that you need to adopt the defaults to your needs.                         !!")
    print("!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!")
    tomlfile = TOMLFile(configurationFilename)
    config = tomlkit.parse(defaultConfiguration)
    tomlfile.write(config)


def loadConfiguration(configurationFilename):
    """Load the configuration from a persitant file. 
    If the configuration file does not exist, then a default file is written, which is intended as an example.

    Parameters
    ----------
    configurationFilename : string
        The configuration filename.

    Returns
    -------
    myConfig :  dictionary
        The configuration dictionary.
    """
    
    if not os.path.isfile(configurationFilename):
        generateDefaultConfigurationTomlFile(configurationFilename)
   
    myConfig = tomlkit.toml_file.TOMLFile(configurationFilename).read()

    return myConfig


def printAvailableWorksheetsIn(sourceDataFilename, xlsxWorkbook, contentSheetname):
    print("Available worksheets in '", sourceDataFilename, "' are: ", sep="")
    for sheetname in xlsxWorkbook.sheetnames:
        sheet = xlsxWorkbook[sheetname]
        print(sheetname, "\tmax_row:", sheet.max_row, "\tmax_column:", sheet.max_column, end="")

        if sheetname == contentSheetname:
            print(" Will be used to generate output files.")
        else:
            print("")


def xlx2docx(myConfig):
    sourceDataFilename = myConfig["sourceData"    ]["filename"         ]
    contentSheetname   = myConfig["sourceData"    ]["contentSheetname" ]
    contentStartRow    = myConfig["sourceData"    ]["contentStartRow"  ]
    commandColumn      = myConfig["sourceData"    ]["commandColumn"    ]
    styleColumn        = myConfig["sourceData"    ]["styleColumn"      ]
    replaceColumn      = myConfig["sourceData"    ]["replaceColumn"    ]
    contentColumns     = myConfig["sourceData"    ]["contentColumns"   ]
    templateFilename   = myConfig["template"      ]["filename"         ]
    outputFilenames    = myConfig["generatedData" ]["filenames"        ]

    print("Configuration:\n",
          "Source data file          '" , sourceDataFilename ,     "'\n" ,
          "From sheet                '" , contentSheetname   ,     "'\n" ,
          "starting at row           "  , contentStartRow    ,     "\n"  ,
          "using command from column "  , commandColumn      ,     "\n"  ,
          "using style from column   "  , styleColumn        ,     "\n"  ,
          "using replace from column "  , replaceColumn      ,     "\n"  ,
          "using content columns     "  , contentColumns     ,     "\n"  ,
          "using template file       '" , templateFilename   ,     "'\n" ,
          "to generate files         "  , outputFilenames    , sep=""    ,
          )

    #######################################################################################
    # Load xlsx source data.
    #######################################################################################
    # If data_only is true, then the formulas are replaced by value. 
    # The values are not evaluated by docx. 
    # The values were generated by MS-Excel at the time of last edit and stored within the file.
    xlsxWorkbook = openpyxl.load_workbook(sourceDataFilename, data_only=True) 

    printAvailableWorksheetsIn(sourceDataFilename, xlsxWorkbook, contentSheetname)

    #######################################################################################
    # Generate MS-Word files from MS-Excel columns.
    #######################################################################################
    for sheetname in xlsxWorkbook.sheetnames:
        if sheetname == contentSheetname:
            sheet = xlsxWorkbook[contentSheetname]
            generateDocxFromXlsx(
                sheet            = sheet            ,
                contentStartRow  = contentStartRow  ,
                commandColumn    = commandColumn    ,
                styleColumn      = styleColumn      ,
                replaceColumn    = replaceColumn    ,
                contentColumns   = contentColumns   ,
                templateFilename = templateFilename ,
                outputFilenames  = outputFilenames  ,
                )
            break


def printUsage():
    print(__file__, __version__)
    usage = f"""
    --version : prints version
    --help    : prints this help message
    
    An alternative configurationFileName can be passed as argument. Example: '{__file__} somewhereElse/myAlternativeConfigfile.toml'
    """

    print(usage)



def main():
    configurationFilename = "xlsx2docx-byColumns_Configuration.toml"

    if len(sys.argv) > 1:
        if sys.argv[1] == "--version":
            print(f"{__file__} V{__version__}")
            sys.exit()
        elif sys.argv[1] == "--help":
            printUsage()
            sys.exit()
        else:
            print(f"Use of alternative configuration file '{sys.argv[1]}'")
            configurationFilename = sys.argv[1]

    myConfig = loadConfiguration(configurationFilename)

    xlx2docx(myConfig)


if __name__ == "__main__":
    # execute only if run as a script
    main()
